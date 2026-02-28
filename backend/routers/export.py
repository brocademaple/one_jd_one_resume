import io
import os
import re
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer
from sqlalchemy.orm import Session

from database import get_db
import models

router = APIRouter(prefix="/api/export", tags=["export"])


def _clean_content(content: str) -> str:
    content = re.sub(r"===RESUME_START===\n?", "", content)
    return re.sub(r"===RESUME_END===\n?", "", content)


def _safe_filename(title: Optional[str], suffix: str, fallback: str) -> str:
    if title:
        # Keep letters, digits, spaces, hyphens, and CJK (Unicode word chars)
        safe_title = re.sub(r'[<>:"/\\|?*]', "", title).strip()[:50]
        if safe_title:
            return f"{safe_title}.{suffix}"
    return fallback


def _strip_md_inline(text: str) -> str:
    """Remove markdown bold/italic/code from text for Word export."""
    s = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    s = re.sub(r"\*(.*?)\*", r"\1", s)
    s = re.sub(r"_(.*?)_", r"\1", s)
    s = re.sub(r"__(.*?)__", r"\1", s)
    s = re.sub(r"`(.*?)`", r"\1", s)
    return s


def md_to_pdf_content(md_text: str, font_size: int = 10):
    """font_size: 9, 10, 11, 12 (body text base)."""
    font_path_regular = None
    font_candidates = [
        # Windows
        os.path.expandvars(r"%WINDIR%\Fonts\msyh.ttc"),
        os.path.expandvars(r"%WINDIR%\Fonts\msyhbd.ttc"),
        os.path.expandvars(r"%WINDIR%\Fonts\simsun.ttc"),
        "C:\\Windows\\Fonts\\msyh.ttc",
        "C:\\Windows\\Fonts\\simsun.ttc",
        # macOS
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
        # Linux
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    ]
    for fp in font_candidates:
        if not fp:
            continue
        normalized = os.path.normpath(os.path.expanduser(os.path.expandvars(fp)))
        if os.path.exists(normalized):
            font_path_regular = normalized
            break

    if font_path_regular:
        try:
            pdfmetrics.registerFont(TTFont("ChineseFont", font_path_regular))
            base_font = "ChineseFont"
        except Exception:
            base_font = "Helvetica"
    else:
        base_font = "Helvetica"

    getSampleStyleSheet()
    h1_size = max(14, font_size + 8)
    h2_size = max(11, font_size + 3)
    h3_size = max(10, font_size + 1)
    style_h1 = ParagraphStyle("H1", fontName=base_font, fontSize=h1_size, spaceAfter=6, spaceBefore=10, textColor=colors.HexColor("#1a1a2e"), leading=h1_size + 4)
    style_h2 = ParagraphStyle("H2", fontName=base_font, fontSize=h2_size, spaceAfter=4, spaceBefore=8, textColor=colors.HexColor("#16213e"), leading=h2_size + 4)
    style_h3 = ParagraphStyle("H3", fontName=base_font, fontSize=h3_size, spaceAfter=3, spaceBefore=5, textColor=colors.HexColor("#0f3460"), leading=h3_size + 3)
    style_body = ParagraphStyle("Body", fontName=base_font, fontSize=font_size, spaceAfter=2, spaceBefore=1, leading=font_size + 4, textColor=colors.HexColor("#333333"))
    style_bullet = ParagraphStyle("Bullet", fontName=base_font, fontSize=font_size, spaceAfter=2, spaceBefore=1, leading=font_size + 4, leftIndent=15, textColor=colors.HexColor("#333333"))

    story = []
    for line in md_text.split("\n"):
        stripped = re.sub(r"\*\*(.*?)\*\*", r"\1", line.rstrip())
        stripped = re.sub(r"\*(.*?)\*", r"\1", stripped)
        stripped = re.sub(r"`(.*?)`", r"\1", stripped)

        if line.startswith("# "):
            story.append(Paragraph(stripped[2:].strip(), style_h1))
        elif line.startswith("## "):
            story.append(Spacer(1, 0.1 * cm))
            story.append(Paragraph(stripped[3:].strip(), style_h2))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cccccc")))
        elif line.startswith("### "):
            story.append(Paragraph(stripped[4:].strip(), style_h3))
        elif line.startswith("- ") or line.startswith("* "):
            text = stripped[2:].strip()
            if text:
                story.append(Paragraph(f"• {text}", style_bullet))
        elif line.strip() in ["---", "***"]:
            story.append(HRFlowable(width="100%", thickness=0.3, color=colors.HexColor("#dddddd")))
        elif not line.strip():
            story.append(Spacer(1, 0.2 * cm))
        elif stripped.strip():
            story.append(Paragraph(stripped, style_body))

    return story


@router.get("/pdf/{resume_id}")
def export_pdf(
    resume_id: int,
    db: Session = Depends(get_db),
    font_size: int = 10,
    margin_cm: float = 2.0,
):
    """font_size: 9-12, margin_cm: 1.5, 2.0, 2.5"""
    font_size = max(9, min(12, font_size))
    margin_cm = max(1.0, min(3.0, margin_cm))
    db_resume = db.query(models.Resume).filter(models.Resume.id == resume_id).first()
    if not db_resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    content = _clean_content(db_resume.content)
    buffer = io.BytesIO()
    m = margin_cm * cm
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=m, leftMargin=m, topMargin=m, bottomMargin=m)
    doc.build(md_to_pdf_content(content, font_size))
    buffer.seek(0)

    filename = _safe_filename(db_resume.title, "pdf", f"resume_{resume_id}.pdf")
    return StreamingResponse(buffer, media_type="application/pdf", headers={"Content-Disposition": f'attachment; filename="{filename}"'})


@router.get("/pdf-preview/{resume_id}")
def preview_pdf(
    resume_id: int,
    db: Session = Depends(get_db),
    font_size: int = 10,
    margin_cm: float = 2.0,
):
    response = export_pdf(resume_id, db, font_size, margin_cm)
    filename = response.headers["Content-Disposition"].split('filename="')[-1].rstrip('"')
    response.headers["Content-Disposition"] = f'inline; filename="{filename}"'
    return response


def _add_paragraph_with_font(doc: Document, text: str, style_name: Optional[str] = None, font_pt: int = 11):
    p = doc.add_paragraph(text, style=style_name)
    for run in p.runs:
        run.font.size = Pt(font_pt)
    return p


@router.get("/word/{resume_id}")
def export_word(
    resume_id: int,
    db: Session = Depends(get_db),
    font_size: int = 11,
    margin_cm: float = 2.0,
):
    """font_size: 10, 11, 12, 14 (pt). margin_cm: 1.5, 2.0, 2.5"""
    font_size = max(9, min(14, font_size))
    margin_cm = max(1.0, min(3.0, margin_cm))
    db_resume = db.query(models.Resume).filter(models.Resume.id == resume_id).first()
    if not db_resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    content = _clean_content(db_resume.content)
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin = Cm(margin_cm)
        section.right_margin = Cm(margin_cm)
    for line in content.splitlines():
        text = line.strip()
        if not text:
            doc.add_paragraph("")
            continue
        clean = _strip_md_inline(text)
        h1_pt = min(22, font_size + 8)
        h2_pt = min(16, font_size + 4)
        h3_pt = min(14, font_size + 2)
        if text.startswith("# "):
            p = doc.add_heading(clean[2:].strip(), level=1)
            for run in p.runs:
                run.font.size = Pt(h1_pt)
        elif text.startswith("## "):
            p = doc.add_heading(clean[3:].strip(), level=2)
            for run in p.runs:
                run.font.size = Pt(h2_pt)
        elif text.startswith("### "):
            p = doc.add_heading(clean[4:].strip(), level=3)
            for run in p.runs:
                run.font.size = Pt(h3_pt)
        elif text.startswith("- ") or text.startswith("* "):
            _add_paragraph_with_font(doc, clean[2:].strip(), style="List Bullet", font_pt=font_size)
        else:
            _add_paragraph_with_font(doc, clean, font_pt=font_size)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = _safe_filename(db_resume.title, "docx", f"resume_{resume_id}.docx")
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
